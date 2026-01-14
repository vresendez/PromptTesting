import os
import docx

DATA_DIR = "data"

# Mapping: Subtype -> Filename
REFERENCE_FILES = {
    "Geneticist": "UserStoriesGenetistsBasedOnFocusGroup.docx",
    "Pediatrician": "UserStoriesPediatriciansBasedOnFocusGroup.docx",
    "Statistician": "UserStoriesStatisticianBasedOnInterview.docx",
    "standard": "UserStoriesParsing.docx" # Default/Baseline
}

def extract_stories_from_docx(filepath):
    """
    Robustly extracts lines from a docx that look like user stories.
    Assumes stories are non-empty paragraphs or table rows.
    """
    stories = []
    try:
        doc = docx.Document(filepath)
        
        # 1. Paragraphs (bullets)
        for para in doc.paragraphs:
            text = para.text.strip()
            # STRICT FILTER: Check for "As a" or "As an"
            lower = text.lower()
            if len(text) > 10 and ("as a " in lower or "as an " in lower): 
                stories.append(text)
                
        # 2. Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    lower = text.lower()
                    if len(text) > 10 and ("as a " in lower or "as an " in lower):
                        stories.append(text)
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        
    return stories

def load_reference_stories(subtype="standard"):
    """
    Returns a list of reference stories for the given persona/subtype.
    """
    # Normalize subtype
    if subtype not in REFERENCE_FILES:
        # Check partial match? Or default to standard
        filename = REFERENCE_FILES["standard"]
    else:
        filename = REFERENCE_FILES[subtype]
        
    path = os.path.join(DATA_DIR, filename)
    
    if not os.path.exists(path):
        print(f"Warning: Reference file {path} not found. Using empty list.")
        return []
        
    return extract_stories_from_docx(path)
